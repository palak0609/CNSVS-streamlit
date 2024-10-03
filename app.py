import streamlit as st
import pdfplumber
import pandas as pd
from docx import Document
from docx.shared import Pt
import re
from io import BytesIO
from docx.shared import RGBColor


def main():
    st.title("PDF Data Extraction and DOCX Conversion")

    # File upload
    uploaded_file = st.file_uploader("Upload a PDF file", type="pdf")

    if uploaded_file is not None:
        with pdfplumber.open(uploaded_file) as pdf:
            tables = []
            for page in pdf.pages:
                extracted_tables = page.extract_tables()
                tables.extend(extracted_tables)

        # Initialize table-related data variables
        df1 = None
        df2 = None

        # Check if any tables exist
        if tables:
            # Extract and process the upper table (second table in the PDF) if present
            if len(tables) > 1:
                table_data_1 = tables[1]
                cleaned_columns_1 = []
                for i, column_name in enumerate(table_data_1[0]):
                    if column_name is None:
                        cleaned_columns_1.append(f"Unnamed_{i}")
                    elif table_data_1[0].count(column_name) > 1:
                        cleaned_columns_1.append(f"{column_name}_{i}")
                    else:
                        cleaned_columns_1.append(column_name)

                df1 = pd.DataFrame(table_data_1[1:], columns=cleaned_columns_1)
                st.write("Extracted Upper Table:")
                st.dataframe(df1)

                # Process the upper table
                selected_columns_1 = df1.iloc[:, [0, 3]]
                selected_columns_1.columns = ["Domain Scores", "Percentile"]
                selected_columns_1['Percentile'] = selected_columns_1['Percentile'].apply(clean_percentile)
                selected_columns_1.dropna(subset=['Percentile'], inplace=True)
                selected_columns_1['Grade'] = selected_columns_1['Percentile'].apply(grading_system)

                st.write("Processed Data with Grades for Upper Table:")
                st.dataframe(selected_columns_1)

        # Extract data from specified test tables
        test_data = extract_vbm_vsm_finger_tests(uploaded_file)
        if test_data is not None:
            st.write("Extracted Test Data:")
            st.dataframe(pd.DataFrame(test_data))

            # Extract and process the lower table containing Domain, Score, and Severity if present
            lower_table_data = None
            for i, table in enumerate(tables):
                # Check each table for the lower table structure based on known headers
                if len(table) > 0 and len(table[0]) >= 3 and table[0][0] == "Domain" and table[0][1] == "Score" and table[0][2] == "Severity":
                    lower_table_data = table
                    break

            if lower_table_data:
                # Handle rows with extra columns
                cleaned_lower_table_data = [row[:3] for row in lower_table_data[1:] if len(row) >= 3]
                df2 = pd.DataFrame(cleaned_lower_table_data, columns=["Domain", "Score", "Severity"])

                st.write("Extracted Lower Table (Domain, Score, Severity):")
                st.dataframe(df2)

                # Clean and process the lower table
                df2['Score'] = df2['Score'].apply(clean_score)
                df2.dropna(subset=['Score'], inplace=True)
                df2['Grade'] = df2['Severity'].apply(grading_system_2)

                st.write("Processed Data with Grades for Lower Table:")
                st.dataframe(df2)

        else:
            st.error("The uploaded PDF does not contain any tables.")
            return

        # Combine both dataframes if available or process separately
        if df1 is not None and df2 is not None:
            combined_df = pd.concat([selected_columns_1, df2], axis=0, ignore_index=True)
        elif df1 is not None:
            combined_df = selected_columns_1
        elif df2 is not None:
            combined_df = df2
        else:
            combined_df = pd.DataFrame()
            st.error("No table data was extracted.")

        # Extract GAD-7 score
        gad7_score = extract_gad7_score(uploaded_file)
        st.write(gad7_score)

        # Extract PHQ-9 score
        phq9_score = extract_phq9_score(uploaded_file)
        phq9_interpretation = interpret_phq9_score(phq9_score)
        st.write(phq9_interpretation)

        # Convert to DOCX and prepare for download
        if not combined_df.empty or test_data is not None:
            docx_data = csv_to_docx_with_flagging(combined_df, test_data, gad7_score, phq9_interpretation)
            st.download_button(
                label="Download DOCX",
                data=docx_data,
                file_name="extracted_data.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        else:
            st.error("No data to convert to DOCX.")

def clean_percentile(value):
    match = re.search(r'\d+', str(value))
    return int(match.group(0)) if match else None

def grading_system(percentile):
    if percentile > 74:
        return "Above Average"
    elif 25 <= percentile <= 74:
        return "Average"
    elif 9 <= percentile <= 24:
        return "Low Average"
    elif 2 <= percentile <= 8:
        return "Low"
    else:
        return "Very Low"

def clean_score(value):
    match = re.search(r'\d+', str(value))
    return int(match.group(0)) if match else None

def grading_system_2(severity):
    if severity in ["Mild", "Moderate", "Severe"]:
        return "FLAG"
    else:
        return None

def csv_to_docx_with_flagging(df, test_data, gad7_score, phq9_score):
    doc = Document()

    # Adding title with adjusted font size
    title = doc.add_paragraph("CNSVS Metrics with Percentiles, Scores, and Grades")
    title_run = title.runs[0]
    title_run.font.size = Pt(16)
    title_run.bold = True  # Make the text bold
    title_run.font.color.rgb = RGBColor(0, 0, 0)  # Set font color to black

    # Flag to track when to insert the new heading
    npq_heading_added = False

    for _, row in df.iterrows():
        if 'Domain Scores' in row and pd.notna(row['Domain Scores']):
            text = f"{row['Domain Scores']}: {row['Percentile']}, {row['Grade']}"
        elif 'Domain' in row and pd.notna(row['Domain']):
            text = f"{row['Domain']}: {row['Score']}, {row['Severity']}"
            # Check if the current row is part of the NeuroPsych Questionnaire (NPQ)
            if row['Domain'] == "Attention" and not npq_heading_added:
                # Add the new heading for the NPQ
                heading = doc.add_paragraph()
                heading_run = heading.add_run("NeuroPsych Questionnaire (NPQ) SF-45")  # Create the run and add the text
                heading_run.bold = True  # Make the text bold
                heading_run.font.size = Pt(14)  # Adjust the size if needed
                heading_run.font.color.rgb = RGBColor(0, 0, 0)  # Set heading color to black
                npq_heading_added = True

        # Create a new paragraph
        p = doc.add_paragraph(style='ListBullet')

        # Add " | FLAG" if the grade is "Low Average," "Low," or "Very Low" for both tables
        if 'Grade' in row and row['Grade'] in ['Low Average', 'Low', 'Very Low']:
            run = p.add_run(text + " | ")
            run_bold = p.add_run("FLAG")
            run_bold.bold = True
        elif 'Grade' in row and row['Grade'] == "FLAG":
            run = p.add_run(text + " | ")
            run_bold = p.add_run("FLAG")
            run_bold.bold = True
        else:
            p.add_run(text)

    # Add extracted test data to the DOCX document
    if test_data is not None:
        for test_name, metrics in test_data.items():
            # Add the test name as a heading
            test_name_paragraph = doc.add_paragraph(f" {test_name}")
            test_name_paragraph.runs[0].bold = True  # Make the test name bold

            # Add each metric under the test
            for metric, value in metrics.items():
                p = doc.add_paragraph(style='ListBullet')
                if "FLAG" in value:
                    value_part, flag_part = value.split(" | ")
                    run = p.add_run(f"  {metric}: {value_part} | ")
                    run_flag = p.add_run("FLAG")
                    run_flag.bold = True
                else:
                    p.add_run(f"  {metric}: {value}")

    # Add GAD-7 score to the document if available
    if gad7_score:
        gad7_paragraph = doc.add_paragraph(gad7_score)
        gad7_paragraph.runs[0].bold = True

    # Add PHQ-9 score to the document if available
    if phq9_score:
        phq9_paragraph = doc.add_paragraph(phq9_score)
        phq9_paragraph.runs[0].bold = True

    # Convert DOCX to a BytesIO object for download
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio


def extract_vbm_vsm_finger_tests(uploaded_file):
    test_data = {
        "Verbal Memory Test (VBM)": {},
        "Visual Memory Test (VSM)": {},
        "Finger Tapping Test (FTT)": {},
        "Symbol Digit Coding (SDC)": {},
        "Stroop Test (ST)": {},
        "Shifting Attention Test (SAT)": {},
        "Continuous Performance Test (CPT)": {},
        "Perception Of Emotions Test (POET)": {},
        "Reasoning Test (RT)": {},
        "Four Part Continuous Performance Test (FPCPT)": {}
    }

    with pdfplumber.open(uploaded_file) as pdf:
        text = ""
        for page in pdf.pages:
            text += page.extract_text()

        # Function to apply flagging
        def apply_flagging(percentile):
            percentile = int(percentile)
            if percentile > 74:
                return f"{percentile}, Above Average"
            elif 25 <= percentile <= 74:
                return f"{percentile}, Average"
            elif 9 <= percentile <= 24:
                return f"{percentile}, Low Average | FLAG"
            elif 2 <= percentile <= 8:
                return f"{percentile}, Low | FLAG"
            else:
                return f"{percentile}, Very Low | FLAG"
            

        # Extract Verbal Memory Test (VBM) data
        vbm_pattern = re.compile(r"Verbal Memory Test \(VBM\).*?Correct Hits - Immediate \d+ \d+ (\d+).*?Correct Passes - Immediate \d+ \d+ (\d+).*?Correct Hits - Delay \d+ \d+ (\d+).*?Correct Passes - Delay \d+ \d+ (\d+)", re.DOTALL)
        vbm_match = vbm_pattern.search(text)
        if vbm_match:
            test_data["Verbal Memory Test (VBM)"] = {
                "Correct Hits - Immediate": apply_flagging(vbm_match.group(1)),
                "Correct Passes - Immediate": apply_flagging(vbm_match.group(2)),
                "Correct Hits - Delay": apply_flagging(vbm_match.group(3)),
                "Correct Passes - Delay": apply_flagging(vbm_match.group(4))
            }

        # Extract Visual Memory Test (VSM) data
        vsm_pattern = re.compile(r"Visual Memory Test \(VSM\).*?Correct Hits - Immediate \d+ \d+ (\d+).*?Correct Passes - Immediate \d+ \d+ (\d+).*?Correct Hits - Delay \d+ \d+ (\d+).*?Correct Passes - Delay \d+ \d+ (\d+)", re.DOTALL)
        vsm_match = vsm_pattern.search(text)
        if vsm_match:
            test_data["Visual Memory Test (VSM)"] = {
                "Correct Hits - Immediate": apply_flagging(vsm_match.group(1)),
                "Correct Passes - Immediate": apply_flagging(vsm_match.group(2)),
                "Correct Hits - Delay": apply_flagging(vsm_match.group(3)),
                "Correct Passes - Delay": apply_flagging(vsm_match.group(4))
            }

        # Extract Finger Tapping Test (FTT) data
        ftt_pattern = re.compile(r"Finger Tapping Test \(FTT\).*?Right Taps Average \d+ \d+ (\d+).*?Left Taps Average \d+ \d+ (\d+)", re.DOTALL)
        ftt_match = ftt_pattern.search(text)
        if ftt_match:
            test_data["Finger Tapping Test (FTT)"] = {
                "Right Taps Average": apply_flagging(ftt_match.group(1)),
                "Left Taps Average": apply_flagging(ftt_match.group(2))
            }

        # Extract Symbol Digit Coding (SDC) data
        sdc_pattern = re.compile(r"Symbol Digit Coding \(SDC\).*?Correct Responses \d+ \d+ (\d+).*?Errors\* \d+ \d+ (\d+)", re.DOTALL)
        sdc_match = sdc_pattern.search(text)
        if sdc_match:
            test_data["Symbol Digit Coding (SDC)"] = {
                "Correct Responses": apply_flagging(sdc_match.group(1)),
                "Errors*": apply_flagging(sdc_match.group(2))
            }

        # Extract Stroop Test (ST) data
        st_pattern = re.compile(r"Stroop Test \(ST\).*?Simple Reaction Time\* \d+ \d+ (\d+).*?Complex Reaction Time Correct\* \d+ \d+ (\d+).*?Stroop Reaction Time Correct\* \d+ \d+ (\d+).*?Stroop Commission Errors\* \d+ \d+ (\d+)", re.DOTALL)
        st_match = st_pattern.search(text)
        if st_match:
            test_data["Stroop Test (ST)"] = {
                "Simple Reaction Time*": apply_flagging(st_match.group(1)),
                "Complex Reaction Time Correct*": apply_flagging(st_match.group(2)),
                "Stroop Reaction Time Correct*": apply_flagging(st_match.group(3)),
                "Stroop Commission Errors*": apply_flagging(st_match.group(4))
            }

        # Extract Shifting Attention Test (SAT) data
        sat_pattern = re.compile(r"Shifting Attention Test \(SAT\).*?Correct Responses \d+ \d+ (\d+).*?Errors\* \d+ \d+ (\d+).*?Correct Reaction Time\* \d+ \d+ (\d+)", re.DOTALL)
        sat_match = sat_pattern.search(text)
        if sat_match:
            test_data["Shifting Attention Test (SAT)"] = {
                "Correct Responses": apply_flagging(sat_match.group(1)),
                "Errors*": apply_flagging(sat_match.group(2)),
                "Correct Reaction Time*": apply_flagging(sat_match.group(3))
            }

        # Extract Continuous Performance Test (CPT) data
        cpt_pattern = re.compile(r"Continuous Performance Test \(CPT\).*?Correct Responses \d+ \d+ (\d+).*?Omission Errors\* \d+ \d+ (\d+).*?Commission Errors\* \d+ \d+ (\d+).*?Choice Reaction Time Correct\* \d+ \d+ (\d+)", re.DOTALL)
        cpt_match = cpt_pattern.search(text)
        if cpt_match:
            test_data["Continuous Performance Test (CPT)"] = {
                "Correct Responses": apply_flagging(cpt_match.group(1)),
                "Omission Errors*": apply_flagging(cpt_match.group(2)),
                "Commission Errors*": apply_flagging(cpt_match.group(3)),
                "Choice Reaction Time Correct*": apply_flagging(cpt_match.group(4))
            }

        # Extract Perception Of Emotions Test (POET) data
        poet_pattern = re.compile(r"Perception Of Emotions Test \(POET\).*?Correct Responses \d+ \d+ (\d+).*?Average Correct Reaction Time\* \d+ \d+ (\d+).*?Omission Errors\* \d+ \d+ (\d+).*?Commission Errors\* \d+ \d+ (\d+).*?Positive Emotions.*?Correct Hits \d+ \d+ (\d+).*?Reaction Time\* \d+ \d+ (\d+).*?Negative Emotions.*?Correct Hits \d+ \d+ (\d+).*?Reaction Time\* \d+ \d+ (\d+)", re.DOTALL)
        poet_match = poet_pattern.search(text)
        if poet_match:
            test_data["Perception Of Emotions Test (POET)"] = {
                "Correct Responses": apply_flagging(poet_match.group(1)),
                "Average Correct Reaction Time*": apply_flagging(poet_match.group(2)),
                "Omission Errors*": apply_flagging(poet_match.group(3)),
                "Commission Errors*": apply_flagging(poet_match.group(4)),
                "Positive Emotions": [
                    {"Correct Hits": apply_flagging(poet_match.group(5))},
                    {"Reaction Time*": apply_flagging(poet_match.group(6))}
                ],
                "Negative Emotions": [
                    {"Correct Hits": apply_flagging(poet_match.group(7))},
                    {"Reaction Time*": apply_flagging(poet_match.group(8))}
                ]
            }

        # Extract Reasoning Test (RT) data
        rt_pattern = re.compile(r"Reasoning Test \(RT\).*?Correct Responses \d+ \d+ (\d+).*?Average Correct Reaction Time\* \d+ \d+ (\d+).*?Commission Errors\* \d+ \d+ (\d+).*?Omission Errors\* \d+ \d+ (\d+)", re.DOTALL)
        rt_match = rt_pattern.search(text)
        if rt_match:
            test_data["Reasoning Test (RT)"] = {
                "Correct Responses": apply_flagging(rt_match.group(1)),
                "Average Correct Reaction Time*": apply_flagging(rt_match.group(2)),
                "Commission Errors*": apply_flagging(rt_match.group(3)),
                "Omission Errors*": apply_flagging(rt_match.group(4))
            }

        # Extract Four Part Continuous Performance Test (FPCPT) data
        fpcpt_pattern = re.compile(r"Four Part Continuous Performance Test \(FPCPT\).*?"
                        r"Part 1.*?Average Correct Reaction Time\* \d+ \d+ (\d+).*?"
                        r"Part 2.*?Correct Responses \d+ \d+ (\d+).*?Average Correct Reaction Time\* \d+ \d+ (\d+).*?"
                        r"Incorrect Responses\* \d+ \d+ (\d+).*?Average Incorrect Reaction Time\* \d+ \d+ (\d+).*?"
                        r"Omission Errors\* \d+ \d+ (\d+).*?"
                        r"Part 3.*?Correct Responses \d+ \d+ (\d+).*?Average Correct Reaction Time\* \d+ \d+ (\d+).*?"
                        r"Incorrect Responses\* \d+ \d+ (\d+).*?Average Incorrect Reaction Time\* \d+ \d+ (\d+).*?"
                        r"Omission Errors\* \d+ \d+ (\d+).*?"
                        r"Part 4.*?Correct Responses \d+ \d+ (\d+).*?Average Correct Reaction Time\* \d+ \d+ (\d+).*?"
                        r"Incorrect Responses\* \d+ \d+ (\d+).*?Average Incorrect Reaction Time\* \d+ \d+ (\d+).*?"
                        r"Omission Errors\* \d+ \d+ (\d+)", re.DOTALL)       
        fpcpt_match = fpcpt_pattern.search(text)
        if fpcpt_match:
            test_data = f"""
    **Four Part Continuous Performance Test (FPCPT)**:
      • **Part 1**:
        • Average Correct Reaction Time*: {apply_flagging(fpcpt_match.group(1))}
      • **Part 2**:
        • Correct Responses: {apply_flagging(fpcpt_match.group(2))}
        • Average Correct Reaction Time*: {apply_flagging(fpcpt_match.group(3))}
        • Incorrect Responses*: {apply_flagging(fpcpt_match.group(4))}
        • Average Incorrect Reaction Time*: {apply_flagging(fpcpt_match.group(5))}
        • Omission Errors*: {apply_flagging(fpcpt_match.group(6))}
      • **Part 3**:
        • Correct Responses: {apply_flagging(fpcpt_match.group(7))}
        • Average Correct Reaction Time*: {apply_flagging(fpcpt_match.group(8))}
        • Incorrect Responses*: {apply_flagging(fpcpt_match.group(9))}
        • Average Incorrect Reaction Time*: {apply_flagging(fpcpt_match.group(10))}
        • Omission Errors*: {apply_flagging(fpcpt_match.group(11))}
      • **Part 4**:
        • Correct Responses: {apply_flagging(fpcpt_match.group(12))}
        • Average Correct Reaction Time*: {apply_flagging(fpcpt_match.group(13))}
        • Incorrect Responses*: {apply_flagging(fpcpt_match.group(14))}
        • Average Incorrect Reaction Time*: {apply_flagging(fpcpt_match.group(15))}
        • Omission Errors*: {apply_flagging(fpcpt_match.group(16))}
    """
    return test_data

# Function to classify the GAD-7 score based on provided ranges
def classify_gad7_score(score):
    score = int(score)
    if 0 <= score <= 4:
        return "None-Minimal anxiety"
    elif 5 <= score <= 9:
        return "Mild anxiety"
    elif 10 <= score <= 14:
        return "Moderate anxiety"
    elif 15 <= score <= 21:
        return "Severe anxiety"
    else:
        return "Invalid score"

# Function to extract and classify GAD-7 score from a PDF
def extract_gad7_score(pdf_path):
    with pdfplumber.open(pdf_path) as pdf:
        full_text = ""
        for page in pdf.pages:
            full_text += page.extract_text()

    # Regular expression to extract the GAD-7 score and severity
    gad7_pattern = r'GAD-7 Anxiety Severity\s+(\d+)\s*\n'

    # Search for the pattern in the full text
    gad7_match = re.search(gad7_pattern, full_text, re.DOTALL)

    if gad7_match:
        gad7_score = gad7_match.group(1)
        calculated_severity = classify_gad7_score(gad7_score)
        # Format the result to match the required output
        return f"Generalized Anxiety Disorder (GAD-7) Scale:\n• Total Score: {gad7_score} ({calculated_severity})"
    else:
        return "GAD-7 score not found in the document."
    

# Extract PHQ-9 score and classify it
def extract_phq9_score(pdf_path):
    with pdfplumber.open(pdf_path) as pdf:
        text = ""
        for page in pdf.pages:
            text += page.extract_text()
    
    score_match = re.search(r"PHQ-9 Score (\d+)", text)
    
    if score_match:
        score = int(score_match.group(1))
        return score
    else:
        return None

def interpret_phq9_score(score):
    if score is None:
        return "PHQ-9 score not found"
    elif 1 <= score <= 4:
        severity = "Minimal depression"
    elif 5 <= score <= 9:
        severity = "Mild depression"
    elif 10 <= score <= 14:
        severity = "Moderate depression"
    elif 15 <= score <= 19:
        severity = "Moderately severe depression"
    elif 20 <= score <= 27:
        severity = "Severe depression"
    else:
        return "Score out of expected range"

    # Return the output in the desired format
    return f"Patient Health Questionnaire (PHQ-9):\n• Total Score: {score} ({severity})"


if __name__ == "__main__":
    main()
